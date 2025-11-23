class BaseParser:
    @staticmethod
    def parse(file_path):
        raise NotImplementedError


class TextParser(BaseParser):
    @staticmethod
    def parse(file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        metadata = {
            'line_count': text.count('\n') + 1 if text else 0,
            'char_count': len(text),
        }
        return {'text': text, 'metadata': metadata}


class PDFParser(BaseParser):
    @staticmethod
    def parse(file_path):
        try:
            import PyPDF2  # type: ignore
            text = ""
            page_count = 0
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                page_count = len(reader.pages)
                for page in reader.pages:
                    extracted = page.extract_text() or ""
                    text += extracted + "\n"
            metadata = {
                'page_count': page_count,
                'char_count': len(text),
            }
            return {'text': text.strip(), 'metadata': metadata}
        except ImportError:
            raise ImportError("Install PyPDF2")


class DOCXParser(BaseParser):
    @staticmethod
    def parse(file_path):
        try:
            import docx  # type: ignore
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            text = '\n'.join(paragraphs)
            heading_count = sum(1 for p in doc.paragraphs if getattr(p.style, 'name', '').lower().startswith('heading'))
            metadata = {
                'paragraphs': len(paragraphs),
                'heading_count': heading_count,
                'char_count': len(text),
            }
            return {'text': text, 'metadata': metadata}
        except ImportError:
            raise ImportError("Install python-docx")


class CSVParser(BaseParser):
    @staticmethod
    def parse(file_path):
        import csv
        rows = []
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(row)
        text = '\n'.join([' '.join(r) for r in rows])
        metadata = {
            'rows': len(rows),
            'columns': len(rows[0]) if rows else 0,
        }
        return {'text': text, 'metadata': metadata}


class JSONParser(BaseParser):
    @staticmethod
    def parse(file_path):
        import json
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        import json as _json
        text = _json.dumps(data, ensure_ascii=False)
        keys = list(data.keys()) if isinstance(data, dict) else []
        metadata = {
            'type': type(data).__name__,
            'top_level_keys': keys,
        }
        return {'text': text, 'metadata': metadata}


class XMLParser(BaseParser):
    @staticmethod
    def parse(file_path):
        import xml.etree.ElementTree as ET
        tree = ET.parse(file_path)
        root = tree.getroot()

        def _gather_text(node):
            parts = []
            if node.text:
                parts.append(node.text)
            for child in list(node):
                parts.append(_gather_text(child))
            if node.tail:
                parts.append(node.tail)
            return ' '.join([p for p in parts if p])

        text = _gather_text(root)
        tags = set()
        for elem in root.iter():
            tags.add(elem.tag)
        metadata = {
            'root_tag': root.tag,
            'unique_tags': sorted(tags),
            'char_count': len(text),
        }
        return {'text': text.strip(), 'metadata': metadata}


class ImageParser(BaseParser):
    """Parser for image files using OpenAI Vision API."""
    @staticmethod
    def parse(file_path):
        import base64
        from openai import OpenAI
        from django.conf import settings
        import os
        
        # Читаємо зображення і конвертуємо в base64
        with open(file_path, 'rb') as f:
            image_data = f.read()
        
        base64_image = base64.b64encode(image_data).decode('utf-8')
        
        # Визначаємо MIME type на основі розширення
        ext = os.path.splitext(file_path)[1].lower()
        mime_types = {
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.webp': 'image/webp',
        }
        content_type = mime_types.get(ext, 'image/jpeg')
        
        # Використовуємо OpenAI Vision API
        openai_client = OpenAI(api_key=settings.OPENAI_API_KEY)
        
        # Промпт для детального опису зображення
        vision_prompt = "Describe what you see in this image in detail. Include all text, objects, people, and any other relevant information."
        
        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o",  # gpt-4o підтримує vision
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": vision_prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{content_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ],
                max_tokens=1000
            )
            
            text = response.choices[0].message.content
            metadata = {
                'parser': 'image_vision',
                'model': 'gpt-4o',
                'char_count': len(text),
            }
            return {'text': text, 'metadata': metadata}
        except Exception as e:
            # Якщо не вдалося проаналізувати, повертаємо базовий опис
            error_msg = f"Failed to analyze image: {str(e)}"
            return {
                'text': f"[Image file - analysis failed: {error_msg}]",
                'metadata': {'parser': 'image_vision', 'error': str(e)}
            }


def get_parser(file_type):
    parsers = {
        'txt': TextParser,
        'text': TextParser,
        'pdf': PDFParser,
        'docx': DOCXParser,
        'csv': CSVParser,
        'json': JSONParser,
        'xml': XMLParser,
        'jpg': ImageParser,
        'jpeg': ImageParser,
        'png': ImageParser,
        'gif': ImageParser,
        'webp': ImageParser,
        'image': ImageParser,
    }
    return parsers.get(file_type, TextParser)

